# src/agents/case_intake_agent/helpers/evidence/extract_docx_text.py

import asyncio
from docx import Document


async def extract_docx_text(
    file_path: str,
) -> str:
    """
    Extract plain text from a DOCX document.

    Returns:
        Combined text from paragraphs and tables.
    """

    return await asyncio.to_thread(
        _extract_docx_text_sync,
        file_path,
    )


def _extract_docx_text_sync(
    file_path: str,
) -> str:

    document = Document(file_path)

    content = []

    # Paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            content.append(text)

    # Tables
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                content.append(" | ".join(row_text))

    return "\n".join(content)